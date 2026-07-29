from datetime import UTC, datetime
from io import BytesIO

from docx import Document
from docx.document import Document as DocumentObject
from PIL import Image


def representative_docx() -> bytes:
    document = Document()
    properties = document.core_properties
    properties.title = "Extraction fixture"
    properties.author = "MakeTime"
    properties.subject = "R10-S04"
    properties.keywords = "rag,docx"
    properties.created = datetime(2026, 7, 29, 12, 0, tzinfo=UTC)
    properties.modified = datetime(2026, 7, 29, 13, 30, tzinfo=UTC)

    document.add_heading("Document heading", level=2)
    document.add_paragraph("Opening paragraph")

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Service"
    table.cell(0, 1).text = "Responsibility"
    table.cell(1, 0).text = "Python"
    table.cell(1, 1).text = "Extraction"

    document.add_paragraph("Closing paragraph")
    return _save(document)


def merged_table_docx() -> bytes:
    document = Document()
    table = document.add_table(rows=2, cols=2)
    merged = table.cell(0, 0).merge(table.cell(0, 1))
    merged.text = "Merged heading"
    table.cell(1, 0).text = "Left"
    table.cell(1, 1).text = "Right"
    return _save(document)


def nested_table_docx() -> bytes:
    document = Document()
    outer = document.add_table(rows=1, cols=1)
    outer.cell(0, 0).paragraphs[0].text = "Outer text"
    nested = outer.cell(0, 0).add_table(rows=1, cols=2)
    nested.cell(0, 0).text = "Nested"
    nested.cell(0, 1).text = "Values"
    return _save(document)


def empty_docx() -> bytes:
    return _save(Document())


def mixed_text_and_image_docx() -> bytes:
    document = Document()
    document.add_paragraph("Extractable text")
    image_buffer = BytesIO()
    Image.new("RGB", (40, 20), color="black").save(image_buffer, format="PNG")
    image_buffer.seek(0)
    document.add_picture(image_buffer)
    return _save(document)


def _save(document: DocumentObject) -> bytes:
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()

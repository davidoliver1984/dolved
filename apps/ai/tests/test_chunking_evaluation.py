import json
import statistics
from collections import defaultdict
from io import BytesIO
from pathlib import Path
from typing import Any, cast
from uuid import UUID

from docx import Document

from app.chunking import BaselineStructuralChunker, ChunkingResult
from app.extraction import (
    DocxSourceLocation,
    ExtractionContext,
    ExtractorIdentity,
    PdfSourceLocation,
    PlainTextExtractor,
    PythonDocxExtractor,
    TableCell,
    TableRow,
    TextSourceLocation,
)
from app.extraction.models import ExtractedDocumentMetadata
from app.extraction.pdf import PdfPlumberExtractor
from app.normalisation import (
    NormalisedDocument,
    NormalisedTableElement,
    NormaliserIdentity,
    StructuralNormaliser,
)
from tests.pdf_fixtures import PositionedText, text_pdf

WORKSPACE_ID = UUID("79d8699d-2ae1-4072-9ad1-f9dd16f7b325")
DOCUMENT_IDS = {
    "short_plain_text": UUID("10000000-0000-4000-8000-000000000001"),
    "long_plain_text": UUID("10000000-0000-4000-8000-000000000002"),
    "awkward_plain_text": UUID("10000000-0000-4000-8000-000000000003"),
    "pdf": UUID("10000000-0000-4000-8000-000000000004"),
    "docx": UUID("10000000-0000-4000-8000-000000000005"),
    "table": UUID("10000000-0000-4000-8000-000000000006"),
}
CORPUS_PATH = Path(__file__).parent / "fixtures/chunking/corpus.json"


def corpus() -> dict[str, Any]:
    return cast(dict[str, Any], json.loads(CORPUS_PATH.read_text()))


def context(case: str, media_type: str) -> ExtractionContext:
    return ExtractionContext(
        workspace_id=WORKSPACE_ID,
        document_id=DOCUMENT_IDS[case],
        source_media_type=media_type,
    )


def normalised_plain_text(case: str, paragraphs: list[str]) -> NormalisedDocument:
    source = "\n\n".join(paragraphs).encode()
    extracted = PlainTextExtractor().extract(
        source,
        context=context(case, "text/plain"),
    )
    return StructuralNormaliser().normalise(extracted)


def evaluated_cases() -> dict[str, NormalisedDocument]:
    cases = corpus()["cases"]
    long_case = cases["long_plain_text"]
    long_paragraphs = long_case["paragraphs"] * long_case["repeat"]

    return {
        "short_plain_text": normalised_plain_text(
            "short_plain_text", cases["short_plain_text"]["paragraphs"]
        ),
        "long_plain_text": normalised_plain_text("long_plain_text", long_paragraphs),
        "awkward_plain_text": normalised_plain_text(
            "awkward_plain_text", cases["awkward_plain_text"]["paragraphs"]
        ),
        "pdf": normalised_pdf(cases["pdf"]),
        "docx": normalised_docx(cases["docx"]),
        "table": normalised_table(cases["table"]["rows"]),
    }


def normalised_pdf(case: dict[str, Any]) -> NormalisedDocument:
    pages = []
    for page in range(1, case["pages"] + 1):
        entries = [PositionedText(case["repeated_header"], 72, 750)]
        entries.extend(
            PositionedText(
                case["body_template"].format(page=page, paragraph=paragraph),
                72,
                700 - (paragraph * 80),
            )
            for paragraph in range(1, case["paragraphs_per_page"] + 1)
        )
        entries.append(PositionedText(case["repeated_footer"], 72, 40))
        pages.append(entries)
    source = text_pdf(pages, title="R11-S03 prose evaluation")
    extracted = PdfPlumberExtractor().extract(
        source,
        context=context("pdf", "application/pdf"),
    )
    return StructuralNormaliser().normalise(extracted)


def normalised_docx(case: dict[str, Any]) -> NormalisedDocument:
    source_document = Document()
    for section in case["sections"]:
        source_document.add_heading(section["heading"], level=2)
        source_document.add_paragraph(section["paragraph"])
    buffer = BytesIO()
    source_document.save(buffer)
    source = buffer.getvalue()
    extracted = PythonDocxExtractor().extract(
        source,
        context=context(
            "docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ),
    )
    return StructuralNormaliser().normalise(extracted)


def normalised_table(row_count: int) -> NormalisedDocument:
    rows = tuple(
        TableRow(
            index=index,
            cells=(
                TableCell(
                    row_index=index,
                    column_index=0,
                    text=f"service-{index}",
                    source_location=TextSourceLocation(
                        start_character=index,
                        end_character=index + 1,
                        start_line=index + 1,
                        end_line=index + 1,
                    ),
                ),
                TableCell(
                    row_index=index,
                    column_index=1,
                    text=(
                        "Retains deterministic lifecycle, provenance, and "
                        f"retrieval behaviour for evaluation row {index}."
                    ),
                    source_location=TextSourceLocation(
                        start_character=index,
                        end_character=index + 1,
                        start_line=index + 1,
                        end_line=index + 1,
                    ),
                ),
            ),
        )
        for index in range(row_count)
    )
    text = "\n".join("\t".join(cell.text for cell in row.cells) for row in rows)
    element = NormalisedTableElement(
        id=UUID("20000000-0000-4000-8000-000000000001"),
        text=text,
        rows=rows,
        source_element_ids=(UUID("30000000-0000-4000-8000-000000000001"),),
        source_locations=(
            TextSourceLocation(
                start_character=0,
                end_character=len(text),
                start_line=1,
                end_line=row_count,
            ),
        ),
        start_character=0,
        end_character=len(text),
    )
    return NormalisedDocument(
        workspace_id=WORKSPACE_ID,
        document_id=DOCUMENT_IDS["table"],
        source_media_type="text/tab-separated-values",
        source_byte_size=len(text.encode()),
        source_extractor=ExtractorIdentity(name="evaluation-table", version="1"),
        normaliser=NormaliserIdentity(name="structural", version="1"),
        text=text,
        elements=(element,),
        metadata=ExtractedDocumentMetadata(title="R11-S03 table evaluation"),
    )


def primary_text(result: ChunkingResult) -> dict[UUID, str]:
    pieces: dict[UUID, list[tuple[int, str]]] = defaultdict(list)
    for chunk in result.chunks:
        for contribution in chunk.contributions:
            if contribution.role != "primary":
                continue
            pieces[contribution.normalised_element_id].append(
                (
                    contribution.element_start_character,
                    chunk.text[
                        contribution.chunk_start_character : contribution.chunk_end_character
                    ],
                )
            )
    return {
        element_id: "".join(text for _, text in sorted(element_pieces))
        for element_id, element_pieces in pieces.items()
    }


def token_distribution(result: ChunkingResult) -> dict[str, float | int]:
    counts = [chunk.token_count for chunk in result.chunks]
    return {
        "chunks": len(counts),
        "minimum": min(counts),
        "median": statistics.median(counts),
        "mean": round(statistics.mean(counts), 2),
        "maximum": max(counts),
    }


def test_corpus_is_repository_authored_and_documents_expected_boundaries() -> None:
    data = corpus()

    assert (
        data["license"] == "Repository-authored synthetic evaluation material; CC0-1.0."
    )
    assert set(data["cases"]) == set(DOCUMENT_IDS)
    assert all(len(case["expected"]) >= 3 for case in data["cases"].values())


def test_evaluation_corpus_is_deterministic_bounded_and_lossless() -> None:
    chunker = BaselineStructuralChunker()

    for document in evaluated_cases().values():
        first = chunker.chunk(document)
        second = chunker.chunk(document)

        assert first == second
        assert all(chunk.token_count <= 512 for chunk in first.chunks)
        assert primary_text(first) == {
            element.id: element.text for element in document.elements if element.text
        }


def test_provenance_spans_slice_chunk_and_source_content_exactly() -> None:
    for document in evaluated_cases().values():
        result = BaselineStructuralChunker().chunk(document)
        elements = {element.id: element for element in document.elements}

        for chunk in result.chunks:
            for contribution in chunk.contributions:
                element = elements[contribution.normalised_element_id]
                assert (
                    chunk.text[
                        contribution.chunk_start_character : contribution.chunk_end_character
                    ]
                    == element.text[
                        contribution.element_start_character : contribution.element_end_character
                    ]
                )
                assert contribution.source_element_ids == element.source_element_ids
                assert contribution.source_locations == element.source_locations


def test_format_specific_source_locations_survive_to_chunks() -> None:
    cases = evaluated_cases()
    pdf_result = BaselineStructuralChunker().chunk(cases["pdf"])
    docx_result = BaselineStructuralChunker().chunk(cases["docx"])

    assert all(
        isinstance(location, PdfSourceLocation)
        for chunk in pdf_result.chunks
        for contribution in chunk.contributions
        for location in contribution.source_locations
    )
    assert all(
        isinstance(location, DocxSourceLocation)
        for chunk in docx_result.chunks
        for contribution in chunk.contributions
        for location in contribution.source_locations
    )


def test_structural_expectations_hold_for_headings_pages_and_tables() -> None:
    cases = evaluated_cases()
    pdf = cases["pdf"]
    docx = cases["docx"]
    table_result = BaselineStructuralChunker().chunk(cases["table"])

    assert all(
        element.text not in {"MakeTime ingestion handbook", "Internal evaluation copy"}
        for element in pdf.elements
    )
    assert [element.kind for element in docx.elements] == [
        "heading",
        "paragraph",
        "heading",
        "paragraph",
        "heading",
        "paragraph",
    ]
    for heading_index in (0, 2, 4):
        heading_id = docx.elements[heading_index].id
        paragraph_id = docx.elements[heading_index + 1].id
        assert any(
            any(
                primary_ids[index : index + 2] == [heading_id, paragraph_id]
                for index in range(len(primary_ids) - 1)
            )
            for chunk in BaselineStructuralChunker().chunk(docx).chunks
            if (
                primary_ids := [
                    item.normalised_element_id
                    for item in chunk.contributions
                    if item.role == "primary"
                ]
            )
        )
    assert any(
        warning.code == "oversized_table_split" for warning in table_result.warnings
    )
    table_primary_fragments = [
        chunk.text[
            contribution.chunk_start_character : contribution.chunk_end_character
        ]
        for chunk in table_result.chunks
        for contribution in chunk.contributions
        if contribution.role == "primary"
    ]
    assert all(
        fragment.endswith("\n") or fragment == table_primary_fragments[-1]
        for fragment in table_primary_fragments
    )


def test_chunk_size_distributions_are_inspectable(capsys: Any) -> None:
    distributions = {
        name: token_distribution(BaselineStructuralChunker().chunk(document))
        for name, document in evaluated_cases().items()
    }
    print(json.dumps(distributions, indent=2, sort_keys=True))

    assert distributions["short_plain_text"]["chunks"] == 1
    assert distributions["long_plain_text"]["chunks"] > 1
    assert distributions["table"]["chunks"] > 1
    assert capsys.readouterr().out.startswith("{")
